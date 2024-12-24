import json
from docx import Document

def extract_document_info(filepath, output_json="document_info.json"):
    doc = Document(filepath)

    document_info = {
        "标题信息": [],
        "段落信息": [],
        "文档属性": {},
        "页眉": [],
        "页脚": [],
    }

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            document_info["标题信息"].append({
                "标题": paragraph.text.strip(),
                "级别": paragraph.style.name,
                "大纲级别": paragraph.style.priority if hasattr(paragraph.style, "priority") else None
            })

    for paragraph in doc.paragraphs:
        document_info["段落信息"].append({
            "内容": paragraph.text.strip(),
            "样式": paragraph.style.name,
            "缩进": {
                "左缩进": paragraph.paragraph_format.left_indent.pt if paragraph.paragraph_format.left_indent else 0,
                "右缩进": paragraph.paragraph_format.right_indent.pt if paragraph.paragraph_format.right_indent else 0,
                "首行缩进": paragraph.paragraph_format.first_line_indent.pt if paragraph.paragraph_format.first_line_indent else 0,
            },
            "对齐方式": paragraph.alignment if paragraph.alignment else "默认",
        })

    core_props = doc.core_properties
    document_info["文档属性"] = {
        "标题": core_props.title,
        "作者": core_props.author,
        "主题": core_props.subject,
        "关键词": core_props.keywords,
        "备注": core_props.comments,
        "创建时间": core_props.created.isoformat() if core_props.created else None,
        "修改时间": core_props.modified.isoformat() if core_props.modified else None,
        "类别": core_props.category,
    }

    for section in doc.sections:
        if section.header.is_linked_to_previous:
            document_info["页眉"].append(section.header.paragraphs[0].text if section.header.paragraphs else "")
        if section.footer.is_linked_to_previous:
            document_info["页脚"].append(section.footer.paragraphs[0].text if section.footer.paragraphs else "")

    with open(output_json, "w", encoding="utf-8") as json_file:
        json.dump(document_info, json_file, ensure_ascii=False, indent=4)

    print(f"文档信息已提取并保存到 {output_json}")
    return document_info

filepath = "example.docx"
output_json = "document_info.json"
document_info = extract_document_info(filepath, output_json)

print(json.dumps(document_info, ensure_ascii=False, indent=4))
