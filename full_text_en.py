import json
from docx import Document

def extract_document_info(filepath, output_json="document_info.json"):
    doc = Document(filepath)

    document_info = {
        "headings": [],
        "paragraphs": [],
        "document_properties": {},
        "headers": [],
        "footers": [],
    }

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            document_info["headings"].append({
                "text": paragraph.text.strip(),
                "level": paragraph.style.name,
                "outline_level": paragraph.style.priority if hasattr(paragraph.style, "priority") else None
            })

    for paragraph in doc.paragraphs:
        document_info["paragraphs"].append({
            "content": paragraph.text.strip(),
            "style": paragraph.style.name,
            "indentation": {
                "left_indent": paragraph.paragraph_format.left_indent.pt if paragraph.paragraph_format.left_indent else 0,
                "right_indent": paragraph.paragraph_format.right_indent.pt if paragraph.paragraph_format.right_indent else 0,
                "first_line_indent": paragraph.paragraph_format.first_line_indent.pt if paragraph.paragraph_format.first_line_indent else 0,
            },
            "alignment": str(paragraph.alignment) if paragraph.alignment else "default",
        })

    core_props = doc.core_properties
    document_info["document_properties"] = {
        "title": core_props.title,
        "author": core_props.author,
        "subject": core_props.subject,
        "keywords": core_props.keywords,
        "comments": core_props.comments,
        "created": core_props.created.isoformat() if core_props.created else None,
        "modified": core_props.modified.isoformat() if core_props.modified else None,
        "category": core_props.category,
    }

    for section in doc.sections:
        if section.header.is_linked_to_previous:
            document_info["headers"].append(section.header.paragraphs[0].text if section.header.paragraphs else "")
        if section.footer.is_linked_to_previous:
            document_info["footers"].append(section.footer.paragraphs[0].text if section.footer.paragraphs else "")

    with open(output_json, "w", encoding="utf-8") as json_file:
        json.dump(document_info, json_file, ensure_ascii=False, indent=4)

    print(f"Document information has been extracted and saved to {output_json}")
    return document_info

filepath = "example.docx"
output_json = "document_info.json"
document_info = extract_document_info(filepath, output_json)

print(json.dumps(document_info, ensure_ascii=False, indent=4))
